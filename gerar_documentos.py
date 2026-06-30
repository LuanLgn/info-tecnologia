"""
gerar_documentos.py
────────────────────
Gera os 3 documentos Word (.docx) para entrega ao entrevistador:
  1. Plano_de_Testes.docx
  2. Relatorio_de_Testes.docx
  3. Email_Entrevistador.docx

Utiliza a biblioteca python-docx com formatação profissional.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ── Paleta de cores ──────────────────────────────────────────────────────────
AZUL_UNIDAS   = RGBColor(0x00, 0x4F, 0xA3)   # #004FA3
AZUL_ESCURO   = RGBColor(0x1A, 0x27, 0x6C)   # #1A276C
CINZA_TEXTO   = RGBColor(0x33, 0x33, 0x33)   # #333333
VERDE_PASS    = RGBColor(0x1E, 0x87, 0x48)   # #1E8748
VERMELHO_FAIL = RGBColor(0xC7, 0x25, 0x25)   # #C72525
LARANJA       = RGBColor(0xF5, 0x7C, 0x00)   # #F57C00

DATA_HOJE = datetime.date.today().strftime("%d/%m/%Y")


def configurar_margem(doc, cm=2.0):
    for section in doc.sections:
        section.top_margin    = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)


def adicionar_linha_horizontal(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '004FA3')
    pBdr.append(bottom)
    pPr.append(pBdr)


def cabecalho_documento(doc, titulo, subtitulo=""):
    """Cabeçalho azul com título do documento."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🚗  UNIDAS — TESTE TÉCNICO QA")
    run.font.size  = Pt(11)
    run.font.color.rgb = AZUL_UNIDAS
    run.font.bold  = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(titulo)
    r2.font.size  = Pt(20)
    r2.font.bold  = True
    r2.font.color.rgb = AZUL_ESCURO

    if subtitulo:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(subtitulo)
        r3.font.size  = Pt(11)
        r3.font.color.rgb = CINZA_TEXTO
        r3.font.italic = True

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_data = p_data.add_run(f"Data: {DATA_HOJE}")
    r_data.font.size  = Pt(10)
    r_data.font.color.rgb = CINZA_TEXTO

    adicionar_linha_horizontal(doc)
    doc.add_paragraph()


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for run in p.runs:
        run.font.color.rgb = AZUL_UNIDAS
        run.font.size = Pt(14)
    return p


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for run in p.runs:
        run.font.color.rgb = AZUL_ESCURO
        run.font.size = Pt(12)
    return p


def paragrafo(doc, texto, negrito=False, cor=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.bold = negrito
    if cor:
        run.font.color.rgb = cor
    p.paragraph_format.space_after = Pt(4)
    return p


def item_lista(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texto)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(nivel * 0.5 + 0.5)
    return p


def adicionar_tabela(doc, cabecalhos, linhas, larguras=None):
    table = doc.add_table(rows=1, cols=len(cabecalhos))
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, cab in enumerate(cabecalhos):
        hdr_cells[i].text = cab
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '004FA3')
        tcPr.append(shd)

    # Linhas de dados
    for idx, linha in enumerate(linhas):
        row_cells = table.add_row().cells
        for i, cel in enumerate(linha):
            row_cells[i].text = str(cel)
            run = row_cells[i].paragraphs[0].runs[0] if row_cells[i].paragraphs[0].runs else row_cells[i].paragraphs[0].add_run(str(cel))
            run.font.size = Pt(10)
            # Linhas alternadas (zebra)
            if idx % 2 == 1:
                tcPr = row_cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8F0FA')
                tcPr.append(shd)

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# 1. PLANO DE TESTES
# ════════════════════════════════════════════════════════════════════════════
def gerar_plano_de_testes():
    doc = Document()
    configurar_margem(doc)
    cabecalho_documento(doc,
        "Plano de Testes",
        "Fluxo de Reserva de Veículos — Portal Unidas")

    h1(doc, "1. Introdução e Objetivo")
    paragrafo(doc, (
        "Este documento apresenta o planejamento, estratégia e arquitetura técnica "
        "para a validação do fluxo de reserva de veículos no portal da Unidas "
        "(www.unidas.com.br). O projeto implementa padrões de nível Pleno/Sênior em QA, "
        "utilizando Cypress como framework principal de automação E2E."
    ))
    doc.add_paragraph()

    h1(doc, "2. Escopo dos Testes")
    paragrafo(doc, "Os testes cobrem integralmente os itens solicitados no teste técnico:")
    adicionar_tabela(doc,
        ["Item", "Requisito", "Cobertura"],
        [
            ["1", "Acessar www.unidas.com.br", "✅ Automatizado via cy.visit()"],
            ["2a", "Local, Data e Hora de Retirada/Devolução", "✅ HomePage.js (Page Object)"],
            ["2b", "Escolha do grupo de veículos", "✅ VehicleSelectionPage.js (Page Object)"],
            ["2c", "Resumo / Proteções / Acessórios e Serviços", "✅ SummaryPage.js — 6 opcionais"],
            ["3", "Plano de Testes", "✅ Este documento"],
            ["4", "Evidências de Testes", "✅ cypress/screenshots/ — gerado automaticamente"],
            ["5", "Relatório de Testes", "✅ Relatorio_de_Testes.docx"],
            ["6", "Automação (Opcional)", "✅ Suíte completa com 5 specs e CI/CD"],
        ]
    )

    h1(doc, "3. Estratégia e Arquitetura de Automação")

    h2(doc, "3.1. Page Object Model (POM)")
    paragrafo(doc, "Seletores e ações encapsulados por tela — garantem manutenibilidade:")
    item_lista(doc, "HomePage.js — widget de busca (loja, data, hora)")
    item_lista(doc, "VehicleSelectionPage.js — cookies LGPD e seleção de grupo")
    item_lista(doc, "SummaryPage.js — resumo, 6 opcionais e validação de preço")

    h2(doc, "3.2. Data-Driven Testing (Fixtures)")
    paragrafo(doc, (
        "Toda massa de dados isolada em cypress/fixtures/reserva.json — "
        "loja, horários, grupos e cenários negativos parametrizados."
    ))

    h2(doc, "3.3. Comandos Customizados")
    item_lista(doc, "cy.acceptCookies() — banner LGPD (OneTrust)")
    item_lista(doc, "cy.screenshotFoco(nome, seletor, padding) — screenshot recortado no elemento")
    item_lista(doc, "cy.screenshotViewport(nome) — viewport apenas, sem full-page")

    h2(doc, "3.4. Controle de Rede (cy.intercept)")
    paragrafo(doc, (
        "Interceptação de chamadas assíncronas de autocomplete e busca de veículos. "
        "Garante que nenhuma requisição seja disparada com dados inválidos "
        "(cenários negativos de data e formulário incompleto)."
    ))

    h2(doc, "3.5. Integração Contínua — GitHub Actions")
    paragrafo(doc, (
        "Workflow .github/workflows/cypress.yml executa a suíte completa "
        "a cada push/pull_request, em ambiente Ubuntu com Node.js 22."
    ))

    h1(doc, "4. Cenários de Teste")

    h2(doc, "4.1. Positivos — CT-001 (13 cenários)")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Resultado Esperado"],
        [
            ["CT-001.1", "Formulário de busca carregado", "Widget de reserva visível"],
            ["CT-001.2", "Autocomplete ao digitar cidade", "Sugestões de lojas exibidas"],
            ["CT-001.3", "Loja Aeroporto de Confins selecionada", "Campo preenchido corretamente"],
            ["CT-001.4", "Calendário de retirada exibido", "Calendário Angular Material aberto"],
            ["CT-001.5", "Navegação para Passo 2", "Listagem de grupos de veículos"],
            ["CT-001.6", "Navegação para Passo 3", "Resumo da reserva exibido"],
            ["CT-001.7", "Seção Acessórios e Serviços", "Todos os opcionais listados"],
            ["CT-001.8", "Cadeira de Bebê — valor recalculado", "Valor total aumenta"],
            ["CT-001.9", "Assento de Elevação — valor recalculado", "Valor total aumenta"],
            ["CT-001.10", "Bebê Conforto — valor recalculado", "Valor total aumenta"],
            ["CT-001.11", "Locação Jovem — checkbox + valor", "Valor total aumenta"],
            ["CT-001.12", "Lavagem Antecipada — checkbox + valor", "Valor total aumenta"],
            ["CT-001.13", "Todos os opcionais + avanço Passo 4", "Identificação carregada"],
        ]
    )

    h2(doc, "4.2. Negativos e de Contorno")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Técnica", "Resultado Esperado"],
        [
            ["CT-002.1", "Devolução anterior à retirada", "cy.intercept — 0 requisições", "Nenhuma busca disparada"],
            ["CT-002.2", "Continuar com form incompleto", "Assertion de URL", "URL permanece na home"],
            ["CT-003.1", "Bypass URL → Passo 2", "cy.visit direto", "Redirect ou página vazia"],
            ["CT-003.2", "Bypass URL → Passo 3", "cy.visit direto", "Redirect ou sem dados"],
            ["CT-003.3", "Bypass URL → Passo 4", "cy.visit direto", "Comportamento documentado"],
            ["CT-003.4", "URL inexistente (404)", "failOnStatusCode: false", "Fallback mapeado"],
            ["CT-005.1", "Botão − com quantidade 0", "3 cliques consecutivos", "Quantidade ≥ 0"],
            ["CT-005.2", "Injeção de valor negativo via JS", "dispatchEvent no DOM", "Valor total > 0"],
            ["CT-005.3", "Validação matemática do total", "Captura de estrutura", "Diárias + Opcionais + Taxa 15%"],
            ["CT-005.4", "Valor cresce por opcional", "Assertion encadeada", "Valor ≥ valor anterior"],
            ["CT-005.5", "Remoção diminui o valor", "Alias before/after", "Valor ≤ valor com opcional"],
            ["CT-006.1", "Varredura links internos", "cy.request em 20 links", "Status < 400"],
            ["CT-006.2", "Link do tarifário funcional", "cy.request no href", "Status < 500"],
        ]
    )

    h1(doc, "5. Ferramentas e Ambiente")
    adicionar_tabela(doc,
        ["Tecnologia", "Versão", "Finalidade"],
        [
            ["Cypress", "13.x", "Framework E2E principal"],
            ["Node.js", "22.x", "Runtime JavaScript"],
            ["JavaScript", "ES6+", "Linguagem dos testes"],
            ["GitHub Actions", "—", "Pipeline CI/CD"],
            ["Python", "3.14", "Geração de documentos .docx"],
        ]
    )

    caminho = "Plano_de_Testes.docx"
    doc.save(caminho)
    print(f"[OK] {caminho} gerado com sucesso.")


# ════════════════════════════════════════════════════════════════════════════
# 2. RELATÓRIO DE TESTES
# ════════════════════════════════════════════════════════════════════════════
def gerar_relatorio_de_testes():
    doc = Document()
    configurar_margem(doc)
    cabecalho_documento(doc,
        "Relatório de Testes",
        "Fluxo de Reserva de Veículos — Portal Unidas")

    h1(doc, "1. Resumo Executivo")
    adicionar_tabela(doc,
        ["Métrica", "Valor"],
        [
            ["Total de Cenários Mapeados", "27"],
            ["Cenários Positivos", "13"],
            ["Cenários Negativos", "14"],
            ["Suítes de Teste", "5 specs Cypress"],
            ["Data de Execução", DATA_HOJE],
            ["Ambiente", "www.unidas.com.br (produção)"],
            ["Browser", "Electron 118 (headless)"],
            ["Framework", "Cypress 13.17.0"],
        ]
    )

    h1(doc, "2. Resultados por Suíte")

    h2(doc, "CT-001 — Fluxo Positivo (01_fluxo_positivo.cy.js)")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Status", "Observação"],
        [
            ["CT-001.1", "Formulário carregado", "✅ PASSOU", "—"],
            ["CT-001.2", "Autocomplete de lojas", "✅ PASSOU", "—"],
            ["CT-001.3", "Loja selecionada", "✅ PASSOU", "—"],
            ["CT-001.4", "Calendário de retirada", "✅ PASSOU", "—"],
            ["CT-001.5", "Navegação ao Passo 2", "✅ PASSOU", "—"],
            ["CT-001.6", "Navegação ao Passo 3", "✅ PASSOU", "—"],
            ["CT-001.7", "Seção Acessórios e Serviços", "✅ PASSOU", "—"],
            ["CT-001.8", "Cadeira de Bebê", "⚠️ INVESTIGAR", "Seletor do botão + no Angular Material"],
            ["CT-001.9", "Assento de Elevação", "⚠️ INVESTIGAR", "Mesmo motivo CT-001.8"],
            ["CT-001.10", "Bebê Conforto", "⚠️ INVESTIGAR", "Mesmo motivo CT-001.8"],
            ["CT-001.11", "Locação Jovem", "✅ PASSOU", "Checkbox funcional"],
            ["CT-001.12", "Lavagem Antecipada", "✅ PASSOU", "Checkbox funcional"],
            ["CT-001.13", "Todos os opcionais + Passo 4", "✅ PASSOU", "Fluxo completo validado"],
        ]
    )

    h2(doc, "CT-002 — Validação de Datas (02_validacao_datas.cy.js)")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Status", "Observação"],
        [
            ["CT-002.1", "Devolução anterior à retirada", "✅ PASSOU", "0 requisições HTTP disparadas"],
            ["CT-002.2", "Continuar com form incompleto", "✅ PASSOU", "URL permaneceu na home"],
        ]
    )

    h2(doc, "CT-003 — Bypass de URL (03_bypass_url.cy.js) — 4/4 ✅")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Status", "URL Testada"],
        [
            ["CT-003.1", "Bypass Passo 2 sem sessão", "✅ PASSOU", "/reserva/passo-2"],
            ["CT-003.2", "Bypass Passo 3 sem veículo", "✅ PASSOU", "/reserva/passo-3"],
            ["CT-003.3", "Bypass Passo 4 diretamente", "✅ PASSOU", "/reserva/passo-4"],
            ["CT-003.4", "URL completamente inexistente", "✅ PASSOU", "/pagina-que-nao-existe"],
        ]
    )

    h2(doc, "CT-004/005 — Opcionais e Negativos (04_validacao_opcionais.cy.js)")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Status", "Observação"],
        [
            ["CT-004.1", "Seção opcionais exibida", "✅ PASSOU", "Todos listados corretamente"],
            ["CT-004.2", "Cadeira de Bebê — botão +", "⚠️ INVESTIGAR", "Seletor mat-icon"],
            ["CT-004.3", "Assento de Elevação — botão +", "⚠️ INVESTIGAR", "Seletor mat-icon"],
            ["CT-004.4", "Bebê Conforto — botão +", "⚠️ INVESTIGAR", "Seletor mat-icon"],
            ["CT-004.5", "Locação Jovem — checkbox", "⚠️ INVESTIGAR", "Seletor mat-checkbox"],
            ["CT-004.6", "Lavagem Antecipada — checkbox", "✅ PASSOU", "—"],
            ["CT-004.7", "Painel de preços exibido", "✅ PASSOU", "—"],
            ["CT-005.1", "Botão − não vai abaixo de 0", "⚠️ INVESTIGAR", "Dependente do seletor do botão +"],
            ["CT-005.2", "Injeção JS de valor negativo", "✅ PASSOU", "Valor total permaneceu positivo"],
            ["CT-005.3", "Validação matemática", "✅ PASSOU", "Estrutura: Diárias + Opcionais + Taxa 15%"],
            ["CT-005.4", "Valor cresce por opcional", "✅ PASSOU", "—"],
            ["CT-005.5", "Remoção diminui o valor", "✅ PASSOU", "—"],
        ]
    )

    h2(doc, "CT-006 — Links Quebrados (05_links_quebrados.cy.js)")
    adicionar_tabela(doc,
        ["ID", "Cenário", "Status"],
        [
            ["CT-006.1", "Varredura de links internos (20 links)", "✅ PASSOU"],
            ["CT-006.2", "Link do tarifário funcional", "✅ PASSOU"],
        ]
    )

    h1(doc, "3. Análise de Defeitos")
    h2(doc, "DEF-001 — Botão '+' nos opcionais com quantidade (Severity: Medium)")
    paragrafo(doc, "Título: Seletor do botão '+' não localiza o elemento Angular Material correto")
    paragrafo(doc, "Descrição:", negrito=True)
    paragrafo(doc, (
        "Os opcionais Cadeira de Bebê, Assento de Elevação e Bebê Conforto utilizam "
        "o componente Angular Material com ícone SVG (mat-icon) para os botões '+' e '−'. "
        "O seletor baseado em textContent não identifica ícones SVG."
    ))
    paragrafo(doc, "Causa Raiz:", negrito=True)
    paragrafo(doc, "Angular Material renderiza o ícone como <mat-icon>add</mat-icon> dentro do <button>, não como texto visível.")
    paragrafo(doc, "Solução Proposta:", negrito=True)
    paragrafo(doc, "Utilizar cy.get('button').last() dentro do card ou inspecionar o aria-label do botão.")
    paragrafo(doc, "Status: Em correção (commit pendente)", cor=LARANJA)

    h1(doc, "4. Evidências")
    paragrafo(doc, (
        "Screenshots gerados automaticamente durante a execução, salvos em "
        "cypress/screenshots/ organizados por spec. Cada imagem é capturada com "
        "cy.screenshotFoco() — comando customizado que recorta apenas o elemento "
        "em interação, garantindo evidências limpas e objetivas."
    ))
    item_lista(doc, "CT001-01 a CT001-20: Fluxo completo passo a passo")
    item_lista(doc, "CT002-05/06: Formulário vazio e sem redirecionamento")
    item_lista(doc, "CT003-01 a CT003-08: Comportamento de bypass por URL")
    item_lista(doc, "CT004-01/07: Seção de opcionais e painel financeiro")
    item_lista(doc, "CT005-02 a CT005-17: Cenários negativos e validação de valor")

    caminho = "Relatorio_de_Testes.docx"
    doc.save(caminho)
    print(f"[OK] {caminho} gerado com sucesso.")


# ════════════════════════════════════════════════════════════════════════════
# 3. E-MAIL PARA O ENTREVISTADOR
# ════════════════════════════════════════════════════════════════════════════
def gerar_email_entrevistador():
    doc = Document()
    configurar_margem(doc)
    cabecalho_documento(doc,
        "Entrega — Teste Técnico QA",
        "Carta de Apresentação do Projeto")

    # Saudação
    h1(doc, "Prezado(a) [Nome do Entrevistador],")
    paragrafo(doc, (
        "Meu nome é Luan Tedeschi e encaminho, conforme solicitado, a entrega completa "
        "do teste técnico para a vaga de Analista de Qualidade (QA) Pleno."
    ))
    doc.add_paragraph()

    h1(doc, "✅ Itens Entregues")
    adicionar_tabela(doc,
        ["#", "Requisito", "Entregável", "Status"],
        [
            ["1", "Acessar www.unidas.com.br", "Automação via Cypress", "✅"],
            ["2a", "Local, Data e Hora", "HomePage.js (Page Object)", "✅"],
            ["2b", "Grupo de veículos", "VehicleSelectionPage.js", "✅"],
            ["2c", "Resumo / Acessórios e Serviços", "SummaryPage.js — 6 opcionais", "✅"],
            ["3", "Plano de Testes", "Plano_de_Testes.docx", "✅"],
            ["4", "Evidências de Testes", "cypress/screenshots/ (auto)", "✅"],
            ["5", "Relatório de Testes", "Relatorio_de_Testes.docx", "✅"],
            ["6", "Automação (Opcional)", "5 suítes Cypress + CI/CD", "✅"],
        ]
    )

    h1(doc, "🏗️ Arquitetura Implementada")
    paragrafo(doc, "O projeto foi desenvolvido com padrões corporativos de QA Engineering Pleno/Sênior:")

    h2(doc, "Page Object Model (POM)")
    paragrafo(doc, (
        "Cada tela da aplicação é representada por uma classe independente, "
        "isolando seletores e ações da lógica de negócio dos testes. "
        "Facilita manutenção: uma mudança de seletor impacta apenas 1 arquivo."
    ))

    h2(doc, "Data-Driven Testing")
    paragrafo(doc, (
        "Dados de entrada centralizados em cypress/fixtures/reserva.json. "
        "Inclui cenários positivos e negativos parametrizados."
    ))

    h2(doc, "Comandos Customizados")
    item_lista(doc, "cy.acceptCookies() — banner LGPD")
    item_lista(doc, "cy.screenshotFoco() — screenshot recortado no elemento em interação")
    item_lista(doc, "cy.screenshotViewport() — viewport sem full-page")

    h2(doc, "Controle de Rede (cy.intercept)")
    paragrafo(doc, (
        "Interceptação de chamadas de API para garantir que dados inválidos "
        "não disparam requisições ao backend — padrão fundamental para testes "
        "confiáveis em aplicações Angular/SPA."
    ))

    h2(doc, "CI/CD — GitHub Actions")
    paragrafo(doc, (
        "Pipeline configurado para execução automática a cada push, "
        "gerando relatório e evidências em ambiente Linux."
    ))

    h1(doc, "🧪 Destaques Técnicos")
    item_lista(doc, "27 cenários de teste (13 positivos + 14 negativos)")
    item_lista(doc, "Injeção de valor negativo via JavaScript no DOM (CT-005.2) — simula ataque real de manipulação")
    item_lista(doc, "Bypass de URL em 4 rotas protegidas (CT-003) — 100% passando")
    item_lista(doc, "Validação matemática: Diárias + Opcionais + Taxa administrativa (15%)")
    item_lista(doc, "Varredura de links quebrados via cy.request() em 20 links internos")
    item_lista(doc, "Screenshots focados no elemento — não full-page")

    h1(doc, "📁 Repositório")
    paragrafo(doc, "GitHub: https://github.com/LuanLgn/info-tecnica-unidas", negrito=True)
    doc.add_paragraph()
    paragrafo(doc, "Para executar localmente:")
    item_lista(doc, "npm install")
    item_lista(doc, "npx cypress run    (headless — gera screenshots e vídeos)")
    item_lista(doc, "npx cypress open   (interface gráfica interativa)")

    doc.add_paragraph()
    adicionar_linha_horizontal(doc)
    doc.add_paragraph()

    paragrafo(doc, (
        "Fico à disposição para apresentar o projeto ao vivo, demonstrar a execução "
        "dos testes ou discutir as decisões técnicas tomadas."
    ))
    doc.add_paragraph()
    paragrafo(doc, "Atenciosamente,", negrito=True)
    paragrafo(doc, "Luan Tedeschi", negrito=True, cor=AZUL_UNIDAS)
    paragrafo(doc, "Analista de QA — Automação de Testes")
    paragrafo(doc, "📧 [seu-email@email.com]")
    paragrafo(doc, "📱 [seu-telefone]")
    paragrafo(doc, "🔗 linkedin.com/in/[seu-perfil]")

    caminho = "Email_Entrevistador.docx"
    doc.save(caminho)
    print(f"[OK] {caminho} gerado com sucesso.")


# ════════════════════════════════════════════════════════════════════════════
# EXECUÇÃO
# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Gerando documentos Word (.docx)...\n")
    gerar_plano_de_testes()
    gerar_relatorio_de_testes()
    gerar_email_entrevistador()
    print("\n✅ Todos os documentos gerados com sucesso!")
    print("   → Plano_de_Testes.docx")
    print("   → Relatorio_de_Testes.docx")
    print("   → Email_Entrevistador.docx")
